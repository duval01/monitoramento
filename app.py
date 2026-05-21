import streamlit as st
import pandas as pd
import hashlib
import sqlite3
import os
import csv
import io
from datetime import datetime
import smtplib
from email.message import EmailMessage
import gspread
from google.oauth2.service_account import Credentials
from docx import Document
from docx.shared import Pt, RGBColor
from groq import Groq

_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

@st.cache_resource
def get_gspread_client():
    creds = Credentials.from_service_account_info(
        st.secrets["gcp_service_account"], scopes=_SCOPES
    )
    return gspread.Client(auth=creds)

# Configuração da página
st.set_page_config(page_title="Monitoramento AEST - REURB", layout="wide")

# --- 0. AUTENTICAÇÃO ---
USUARIOS = {
    "admin":     "admin",
    "afonso":    "afonso123",
    "pedro":     "pedro123",
    "glenda":    "glenda123",
    "artur":     "artur123",
    "guilherme": "gui123",
    "thania":    "thania123",
    "ryan":      "ryan123",
}

# Usuários com e-mail corporativo (Office365) — demais usam Gmail
EMAILS_CORPORATIVOS = {
    "admin": "luiz.cruz@desenvolvimento.mg.gov.br",
}

def tela_login():
    _, col_centro, _ = st.columns([1, 1, 1])
    with col_centro:
        st.image("AEST Sede (1).png", width=180)
        st.title("🔐 Acesso ao Sistema")
        with st.form("form_login"):
            usuario = st.text_input("Usuário")
            senha   = st.text_input("Senha", type="password")
            if st.form_submit_button("Entrar", type="primary", use_container_width=True):
                if usuario in USUARIOS and USUARIOS[usuario] == senha:
                    st.session_state["logged_in"] = True
                    st.session_state["usuario"]   = usuario
                    st.rerun()
                else:
                    st.error("Usuário ou senha inválidos.")

_hash_url_check = st.query_params.get("hash")
if not _hash_url_check and not st.session_state.get("logged_in"):
    tela_login()
    st.stop()

# --- 1. CONFIGURAÇÃO DE TEMPLATES ---
TEMPLATES = {
    "reurb": {
        "label":          "Minas REURB",
        "sheet_id":       "1UMGNxxWOPASfGx7Q7jesE_vck8ew0c94-ZHby4yEtOA",
        "aba":            "Capa",
        "skiprows":       33,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "mlpc": {
        "label":          "Minas Livre para Crescer",
        "sheet_id":       "105typDO4GrZBNpBcFGyroOvt_4VtYSECGHXmHc3fhP0",
        "aba":            "Capa",
        "skiprows":       32,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "agua_vales_capa": {
        "label":          "Água dos Vales",
        "sheet_id":       "18wGftV1kww1iSGZU-lX2_WyXzRHzrMh46zngzMDR8k0",
        "aba":            "Capa",
        "skiprows":       28,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "agua_vales_barragens": {
        "label":          "Água dos Vales - Barragens",
        "sheet_id":       "18wGftV1kww1iSGZU-lX2_WyXzRHzrMh46zngzMDR8k0",
        "aba":            "Barragens",
        "skiprows":       28,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "coopera": {
        "label":          "Coopera MG",
        "sheet_id":       "1BRSbW8hl2lgWwtqUWtD1had7JfaJd-LHK_HnPnmTFbA",
        "aba":            "Capa",
        "skiprows":       34,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "investimentos": {
        "label":          "Investimentos",
        "sheet_id":       "1T-rgNslZ9GU3wZuPMDHbQPb1l-vj3hCB_Bart-3BCnQ",
        "aba":            "Capa 2026",
        "skiprows":       34,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "rmbh": {
        "label":          "Segurança Hídrica - RMBH",
        "sheet_id":       "1c9lrshPMGZNe-1qagwKbXz9HburEe_tAJoeL5R8n7yo",
        "aba":            "Capa",
        "skiprows":       27,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "mgtech": {
        "label":          "MG Tech",
        "sheet_id":       "1-OuAVSv9qHqstfMCYFeGL6o9GYzzD2GIuLW5oisysYQ",
        "aba":            "Capa replanejamento 2026",
        "skiprows":       33,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
    },
    "vale_litio": {
        "label":          "Vale do Lítio",
        "sheet_id":       "1A_J__cOMZI6MpcvXxYi-ZnIG51QWluDOtuqlC4oXdNc",
        "aba":            "CAPA",
        "skiprows":       35,
        "colunas_pandas": [0, 1, 2, 3, 7, 10, 11, 14, 16, 17, 19, 20],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    17,
        "nivel_tipo":     "espelho",
    },
    "plac_energia": {
        "label":          "PLAC - Energia",
        "sheet_id":       "1vhOoMc9kiva7yyTa_2wV-r2U68pwzi_8n9q33jZ5_JU",
        "aba":            "Cronograma - ENERGIA",
        "skiprows":       1,
        "colunas_pandas": [0, 1, 2, 3, 11, 12, 13, 16, 18, 19, 23, 25],
        "col_xl_origem":  1,
        "col_xl_marco":   12,
        "col_xl_obs":     24,
        "col_xl_data":    19,
        "nivel_tipo":     "flat",
    },
    "plac_industria": {
        "label":          "PLAC - Indústria",
        "sheet_id":       "1vhOoMc9kiva7yyTa_2wV-r2U68pwzi_8n9q33jZ5_JU",
        "aba":            "Cronograma - INDÚSTRIA",
        "skiprows":       1,
        "colunas_pandas": [0, 1, 2, 3, 11, 12, 13, 16, 18, 19, 23, 25],
        "col_xl_origem":  1,
        "col_xl_marco":   12,
        "col_xl_obs":     24,
        "col_xl_data":    19,
        "nivel_tipo":     "flat",
    },
    "plac_transporte": {
        "label":          "PLAC - Transporte",
        "sheet_id":       "1vhOoMc9kiva7yyTa_2wV-r2U68pwzi_8n9q33jZ5_JU",
        "aba":            "Cronograma - TRANSPORTE",
        "skiprows":       1,
        "colunas_pandas": [0, 1, 2, 3, 11, 12, 13, 16, 18, 19, 23, 25],
        "col_xl_origem":  1,
        "col_xl_marco":   12,
        "col_xl_obs":     24,
        "col_xl_data":    19,
        "nivel_tipo":     "flat",
    },
    "plac_des_sust": {
        "label":          "PLAC - Des. Sustentável",
        "sheet_id":       "1vhOoMc9kiva7yyTa_2wV-r2U68pwzi_8n9q33jZ5_JU",
        "aba":            "Cronograma - Des. Sust. e Ação ",
        "skiprows":       1,
        "colunas_pandas": [0, 1, 2, 3, 11, 12, 13, 16, 18, 19, 23, 25],
        "col_xl_origem":  1,
        "col_xl_marco":   12,
        "col_xl_obs":     24,
        "col_xl_data":    19,
        "nivel_tipo":     "flat",
    },
    "dados_reunioes": {
        "label":          "Reunião Gerencial",
        "sheet_id":       "1UdP19eJ-1rNvnuDtZURBqymYLRV4CHVFcwZs_Tw34mY",
        "aba":            "Plano de Ação 2026",
        "skiprows":       1,
        "colunas_pandas": [0, 2, 3, 4, 7, 9, 10, 13, 15, 16, 19, 21],
        "col_xl_origem":  1,
        "col_xl_marco":   8,
        "col_xl_obs":     20,
        "col_xl_data":    24,
        "nivel_tipo":     "flat",
    },
}

# --- 2. GESTÃO DE BASE DE DADOS (SQLite) ---
def init_db():
    conn = sqlite3.connect('monitoramento_reurb.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS respostas
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  hash_id TEXT, marco TEXT, responsavel TEXT,
                  atualizacao TEXT, nova_data TEXT, data_envio TEXT,
                  concluido TEXT, template TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS formularios_respondidos
                 (hash_id TEXT PRIMARY KEY, respondido_por TEXT, data_envio TEXT)''')
    # migrações para bancos existentes
    for col in ["concluido TEXT", "template TEXT"]:
        try:
            c.execute(f"ALTER TABLE respostas ADD COLUMN {col}")
        except sqlite3.OperationalError:
            pass
    conn.commit()
    conn.close()

def salvar_resposta(hash_id, marco, resp, texto, data, concluido="Não", template="reurb"):
    conn = sqlite3.connect('monitoramento_reurb.db')
    c = conn.cursor()
    data_envio = datetime.now().strftime("%d/%m/%Y %H:%M")
    c.execute("INSERT INTO respostas (hash_id, marco, responsavel, atualizacao, nova_data, data_envio, concluido, template) VALUES (?,?,?,?,?,?,?,?)",
              (hash_id, marco, resp, texto, str(data) if data else "", data_envio, concluido, template))
    conn.commit()
    conn.close()

def buscar_respostas():
    conn = sqlite3.connect('monitoramento_reurb.db')
    df = pd.read_sql_query("SELECT id, hash_id, data_envio, responsavel, marco, atualizacao, nova_data, concluido, template FROM respostas ORDER BY id DESC", conn)
    conn.close()
    return df

def marcar_formulario_respondido(lista_hashes, respondido_por):
    conn = sqlite3.connect('monitoramento_reurb.db')
    c = conn.cursor()
    data_envio = datetime.now().strftime("%d/%m/%Y %H:%M")
    for h in lista_hashes:
        c.execute("INSERT OR IGNORE INTO formularios_respondidos (hash_id, respondido_por, data_envio) VALUES (?,?,?)",
                  (h, respondido_por, data_envio))
    conn.commit()
    conn.close()

def verificar_formulario_respondido(lista_hashes):
    conn = sqlite3.connect('monitoramento_reurb.db')
    c = conn.cursor()
    placeholders = ",".join(["?"] * len(lista_hashes))
    c.execute(f"SELECT respondido_por, data_envio FROM formularios_respondidos WHERE hash_id IN ({placeholders}) LIMIT 1", lista_hashes)
    row = c.fetchone()
    conn.close()
    if row:
        return True, row[0], row[1]
    return False, None, None

def deletar_resposta(id_resp):
    conn = sqlite3.connect('monitoramento_reurb.db')
    c = conn.cursor()
    c.execute("DELETE FROM respostas WHERE id = ?", (id_resp,))
    conn.commit()
    conn.close()

init_db()

# --- LOG DE EXECUÇÃO ---
LOG_PATH = "logs_disparos.csv"
LOG_COLUNAS = ["data_hora", "usuario", "destinatarios", "copias", "marcos"]

def registrar_log(usuario, destinatarios, copias, marcos):
    novo_arquivo = not os.path.exists(LOG_PATH)
    with open(LOG_PATH, "a", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames=LOG_COLUNAS)
        if novo_arquivo:
            writer.writeheader()
        writer.writerow({
            "data_hora":    datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
            "usuario":      usuario,
            "destinatarios": "; ".join(destinatarios),
            "copias":       "; ".join(copias) if copias else "",
            "marcos":       "; ".join(marcos),
        })

# --- FUNÇÃO: INJETAR DADOS NO GOOGLE SHEETS ---
def confirmar_na_planilha(hash_id, texto_novo, nova_data_str, template_nome="reurb"):
    cfg = TEMPLATES[template_nome]
    gc = get_gspread_client()
    sh = gc.open_by_key(cfg["sheet_id"])
    ws = sh.worksheet(cfg["aba"])
    data = ws.get_all_values()

    for i in range(cfg["skiprows"], len(data)):
        row_data  = data[i]
        excel_row = i + 1  # gspread usa indexação 1-based

        val_origem = row_data[cfg["col_xl_origem"] - 1] if len(row_data) >= cfg["col_xl_origem"] else ""
        val_marco  = row_data[cfg["col_xl_marco"]  - 1] if len(row_data) >= cfg["col_xl_marco"]  else ""

        if val_marco:
            row_hash = hashlib.sha256(f"{excel_row}{val_origem}{val_marco}".encode()).hexdigest()[:10]
            if row_hash == hash_id:
                obs_idx   = cfg["col_xl_obs"] - 1
                obs_atual = row_data[obs_idx].strip() if len(row_data) > obs_idx else ""
                assinatura = f"[{datetime.now().strftime('%d/%m/%Y')} - Equipe]: {texto_novo}"
                novo_obs  = f"{obs_atual}\n{assinatura}" if obs_atual and obs_atual not in ["None", "-", ""] else assinatura
                ws.update_cell(excel_row, cfg["col_xl_obs"], novo_obs)

                if nova_data_str and nova_data_str not in ["None", ""]:
                    try:
                        data_fmt = datetime.strptime(nova_data_str, "%Y-%m-%d").strftime("%d/%m/%Y")
                        ws.update_cell(excel_row, cfg["col_xl_data"], data_fmt)
                    except ValueError:
                        pass

                return True

    return False

# --- EXPORTAÇÃO DOCX ---
COLUNAS_DOCX = [
    "Marcos Críticos para Realizar as Entregas", "Responsável",
    "Término Planejado (Linha de Base)", "Tendência de Término",
    "Término Real", "Observações",
]

def gerar_docx(df, template_label):
    df_exp = df.copy()
    for c in ["Tendência de Término", "Término Planejado (Linha de Base)", "Término Real"]:
        df_exp[c] = pd.to_datetime(df_exp[c], errors='coerce').dt.strftime('%d/%m/%Y').fillna("")

    doc = Document()
    doc.add_heading(f"Cronograma — {template_label}", level=1)
    doc.add_paragraph(f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')}")

    tabela = doc.add_table(rows=1, cols=len(COLUNAS_DOCX))
    tabela.style = "Table Grid"

    # cabeçalho
    for i, col in enumerate(COLUNAS_DOCX):
        cell = tabela.rows[0].cells[i]
        cell.text = col
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].paragraph_format.alignment = 1  # CENTER

    nivel_rgb = {
        "p1": RGBColor(0x2A, 0x7B, 0xB5),
        "c1": RGBColor(0x7D, 0x7D, 0x7D),
        "c2": RGBColor(0xD9, 0xD9, 0xD9),
    }

    for _, row in df_exp.iterrows():
        cells = tabela.add_row().cells
        for i, col in enumerate(COLUNAS_DOCX):
            cells[i].text = str(row.get(col, ""))
            run = cells[i].paragraphs[0].runs[0] if cells[i].paragraphs[0].runs else cells[i].paragraphs[0].add_run(str(row.get(col, "")))
            run.font.size = Pt(8)
            nivel = row.get("a1", "t1")
            if nivel in nivel_rgb:
                run.font.color.rgb = nivel_rgb[nivel]
            if nivel in ("p1", "c1", "c2"):
                run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# --- 2. PROCESSAMENTO DOS DADOS (Google Sheets) ---
@st.cache_data(ttl=300)
def carregar_e_limpar_dados(template_nome="reurb"):
    cfg = TEMPLATES[template_nome]
    try:
        gc = get_gspread_client()
        sh = gc.open_by_key(cfg["sheet_id"])
        ws = sh.worksheet(cfg["aba"])
        data = ws.get_all_values()

        # Aplica skiprows e converte para DataFrame com todas as colunas
        data = data[cfg["skiprows"]:]
        df = pd.DataFrame(data)

        # Garante colunas suficientes
        max_col = max(cfg["colunas_pandas"])
        for c in range(df.shape[1], max_col + 1):
            df[c] = ""

        df = df.iloc[:, cfg["colunas_pandas"]].copy()
        df.columns = [
            "Origem", "Data_Origem", "Seção", "Subseção",
            "Marcos Críticos para Realizar as Entregas", "Unidade", "Responsável",
            "Término Planejado (Linha de Base)", "Tendência de Término",
            "Término Real", "Observações", "Pontos_Atencao"
        ]

        df['Linha_Excel'] = df.index + cfg["skiprows"] + 1

        df = df[df["Marcos Críticos para Realizar as Entregas"].str.strip() != ""].fillna("")

        def definir_nivel(row):
            nivel_tipo = cfg.get("nivel_tipo", "standard")
            if nivel_tipo == "flat":
                return "t1"
            if nivel_tipo == "espelho":
                secao  = str(row['Seção']).strip()
                subsec = str(row['Subseção']).strip()
                marco  = str(row['Marcos Críticos para Realizar as Entregas']).strip()
                if marco and marco == secao:  return "c1"
                if marco and marco == subsec: return "c2"
                return "t1"
            s, sub = str(row['Seção']).upper(), str(row['Subseção']).upper()
            if s == "P1": return "p1"
            if s == "C1": return "c1"
            if "C2" in sub or "C3" in sub: return "c2"
            return "t1"

        df['a1'] = df.apply(definir_nivel, axis=1)
        df["HASH_ID"] = df.apply(lambda r: hashlib.sha256(
            f"{r['Linha_Excel']}{r['Origem']}{r['Marcos Críticos para Realizar as Entregas']}".encode()
        ).hexdigest()[:10], axis=1)

        return df
    except Exception as e:
        st.error(f"Erro ao ler planilha '{cfg.get('sheet_id', '')}': {e}")
        return pd.DataFrame()

# --- 3. LÓGICA DE ROTEAMENTO ---
query_params = st.query_params
hash_url     = query_params.get("hash")

# Detecta template: URL tem prioridade (formulário externo), senão usa sessão
if hash_url:
    template_ativo = query_params.get("template", "reurb")
else:
    if "template_ativo" not in st.session_state:
        st.session_state["template_ativo"] = "reurb"
    template_ativo = st.session_state["template_ativo"]

df_base = carregar_e_limpar_dados(template_ativo)

if not hash_url:
    # ---------------------------------------------------------
    # TELA DO TÉCNICO (ADMIN)
    # ---------------------------------------------------------
    with st.sidebar:
        st.image("AEST Sede (1).png", width=180)
        st.markdown(f"**Bem vindo(a), {st.session_state.get('usuario', '')}**")

        template_opcoes = list(TEMPLATES.keys())
        template_sel = st.selectbox(
            "📂 Template",
            options=template_opcoes,
            format_func=lambda k: TEMPLATES[k]["label"],
            index=template_opcoes.index(template_ativo),
        )
        if template_sel != template_ativo:
            st.session_state["template_ativo"] = template_sel
            st.rerun()

        st.header("📧 Disparo de E-mail")
        destinatarios = st.text_input("Destinatários", placeholder="email1@ex.com, email2@ex.com")
        copias = st.text_input("Cópias", placeholder="chefe@ex.com, outro@ex.com")

        marcos_disponiveis = df_base[df_base['a1'] == 't1']["Marcos Críticos para Realizar as Entregas"].unique()
        marcos_para_email = st.multiselect("Escolher Marcos para atualizar:", options=marcos_disponiveis)
        anexar_docx = st.checkbox("📄 Anexar tabela completa (.docx) no e-mail")

        st.divider()
        st.title("Filtros de Tabela")
        opcoes_tempo = ["Sem filtro", "15 dias", "30 dias", "45 dias", "60 dias", "90 dias", "120 dias", "Neste mês"]
        dias_filtro = st.radio("Previsão de Término:", options=opcoes_tempo, index=0)
        f_sem_termino = st.checkbox("Apenas atividades sem término real")
        f_ponto_atencao = st.checkbox("Tem ponto de atenção?")

        def limpar_lista(col, df=df_base):
            return sorted([str(x) for x in df[col].unique() if str(x).strip() != ""])

        f_origem = st.multiselect("Origem - Plano de Ação", options=limpar_lista("Origem"))
        f_secao  = st.multiselect("Seção - Cronograma",     options=limpar_lista("Seção"))

        # Responsável: aplica filtros já selecionados para listar apenas nomes relevantes
        df_para_resp = df_base.copy()
        if f_origem: df_para_resp = df_para_resp[df_para_resp["Origem"].isin(f_origem)]
        if f_secao:  df_para_resp = df_para_resp[df_para_resp["Seção"].isin(f_secao)]
        f_resp = st.multiselect("Responsável", options=limpar_lista("Responsável", df_para_resp))

        st.divider()
        if st.button("🚀 Enviar E-mail de Atualização", type="primary", use_container_width=True):
            if not destinatarios:
                st.error("Informe pelo menos um destinatário.")
            else:
                def normalizar_emails(campo):
                    return [e.strip() for e in campo.replace(";", ",").split(",") if e.strip()]

                lista_dest   = normalizar_emails(destinatarios)
                lista_cc     = normalizar_emails(copias) if copias else []
                usuario_atual = st.session_state.get("usuario", "")
                if usuario_atual in EMAILS_CORPORATIVOS:
                    remetente    = EMAILS_CORPORATIVOS[usuario_atual]
                    senha_app    = st.secrets["senha_corp_admin"]
                    smtp_host    = st.secrets.get("smtp_corp_host", "smtp.office365.com")
                    smtp_porta   = int(st.secrets.get("smtp_corp_porta", 587))
                    smtp_usuario = st.secrets.get("smtp_corp_usuario", remetente)
                else:
                    remetente    = st.secrets["gmail_remetente"]
                    senha_app    = st.secrets["gmail_senha_app"]
                    smtp_host    = "smtp.gmail.com"
                    smtp_porta   = 587
                    smtp_usuario = remetente

                if marcos_para_email:
                    hashes     = df_base[df_base["Marcos Críticos para Realizar as Entregas"].isin(marcos_para_email)]["HASH_ID"].tolist()
                    hash_str   = ",".join(hashes)
                    link       = f"https://monitoramento-aest.streamlit.app/?hash={hash_str}&template={template_ativo}"
                    lista_html = "".join([f"<li><b>{m}</b></li>" for m in marcos_para_email])
                    bloco_marcos = f"""
                        <p>Favor atualizar o status dos seguintes marcos críticos:</p>
                        <ul style="background-color:#f5f5f5;padding:15px;border-left:4px solid #2a7bb5;">{lista_html}</ul>
                        <p><a href="{link}" style="color:#2a7bb5;font-weight:bold;">[Clique aqui para acessar o formulário de atualização]</a></p>"""
                else:
                    bloco_marcos = ""

                corpo_html = f"""
                <html><body style="font-family:Arial,sans-serif;color:#333;">
                    <p>Prezados, bom dia!</p>
                    {bloco_marcos}
                </body></html>"""

                try:
                    with st.spinner("Conectando ao servidor de e-mail..."):
                        msg = EmailMessage()
                        msg['Subject'] = f"Monitoramento {TEMPLATES[template_ativo]['label']}: Atualização de Marcos Críticos"
                        msg['From']    = remetente
                        msg['To']      = ", ".join(lista_dest)
                        if lista_cc: msg['Cc'] = ", ".join(lista_cc)
                        msg.set_content("Ative HTML para ler.")
                        msg.add_alternative(corpo_html, subtype='html')
                        if anexar_docx:
                            label_ativo = TEMPLATES[template_ativo]["label"]
                            nome_docx   = f"Cronograma_{label_ativo}_{datetime.now().strftime('%Y%m%d')}.docx"
                            msg.add_attachment(
                                gerar_docx(df_base, label_ativo),
                                maintype="application",
                                subtype="vnd.openxmlformats-officedocument.wordprocessingml.document",
                                filename=nome_docx,
                            )
                        server = smtplib.SMTP(smtp_host, smtp_porta)
                        server.starttls()
                        server.login(smtp_usuario, senha_app)
                        server.send_message(msg)
                        server.quit()
                    registrar_log(st.session_state["usuario"], lista_dest, lista_cc, marcos_para_email)
                    st.success(f"E-mail enviado para: {', '.join(lista_dest)}")
                except Exception as e:
                    st.error(f"Erro ao enviar e-mail: {e}")

    # Lógica de Filtragem
    df_f = df_base.copy()
    if f_origem: df_f = df_f[df_f["Origem"].isin(f_origem)]
    if f_secao: df_f = df_f[df_f["Seção"].isin(f_secao)]
    if f_resp: df_f = df_f[df_f["Responsável"].isin(f_resp)]
    if f_sem_termino: df_f = df_f[(df_f["Término Real"] == "") | (df_f["a1"] != "t1")]
    if f_ponto_atencao: df_f = df_f[(df_f["Pontos_Atencao"] != "") | (df_f["a1"] != "t1")]

    st.title("📑 Painel de Monitoramento AEST")

    st.markdown("""
    <style>
    div.stTabs [role="tablist"] {
        overflow-x: auto;
        flex-wrap: nowrap;
        scrollbar-width: thin;
        scrollbar-color: #888 #f0f0f0;
        padding-bottom: 4px;
    }
    div.stTabs [role="tablist"]::-webkit-scrollbar { height: 6px; }
    div.stTabs [role="tablist"]::-webkit-scrollbar-track { background: #f0f0f0; border-radius: 4px; }
    div.stTabs [role="tablist"]::-webkit-scrollbar-thumb { background: #888; border-radius: 4px; }
    div.stTabs [role="tablist"]::-webkit-scrollbar-thumb:hover { background: #555; }
    </style>
    """, unsafe_allow_html=True)

    # Busca única: usada para notificações nas abas e conteúdo das abas
    todas_respostas = buscar_respostas()

    def _contagem(k):
        if todas_respostas.empty: return 0
        return len(todas_respostas[todas_respostas['template'] == k])

    def _label_aba(k):
        n = _contagem(k)
        label = TEMPLATES[k]['label']
        return f"📥 {label}  🔴 {n}" if n > 0 else f"📥 {label}"

    tab_labels = ["📊 Visão de Cronograma"] + [_label_aba(k) for k in TEMPLATES]
    tabs = st.tabs(tab_labels)
    tab_cron = tabs[0]

    with tab_cron:
        def style_rows(row):
            if row['a1'] == 'p1': return ['background-color: #2a7bb5; color: white; font-weight: bold'] * len(row)
            if row['a1'] == 'c1': return ['background-color: #7d7d7d; color: white; font-weight: bold'] * len(row)
            if row['a1'] == 'c2': return ['background-color: #d9d9d9; font-weight: bold'] * len(row)
            return [''] * len(row)

        colunas_vistas = ["Marcos Críticos para Realizar as Entregas", "Responsável", "Término Planejado (Linha de Base)", "Tendência de Término", "Término Real", "Observações"]

        df_view = df_f.copy()
        for c in ["Tendência de Término", "Término Planejado (Linha de Base)", "Término Real"]:
            df_view[c] = pd.to_datetime(df_view[c], errors='coerce').dt.strftime('%d/%m/%Y').fillna("")

        st.dataframe(df_view[colunas_vistas + ['a1']].style.apply(style_rows, axis=1),
                     use_container_width=True, hide_index=True, column_order=colunas_vistas)

        nome_arquivo = f"Cronograma_{TEMPLATES[template_ativo]['label']}_{datetime.now().strftime('%Y%m%d')}.docx"
        st.download_button(
            label="📄 Exportar tabela como .docx",
            data=gerar_docx(df_f, TEMPLATES[template_ativo]["label"]),
            file_name=nome_arquivo,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    def analisar_com_ia(marco, resp, texto, data):
        try:
            client = Groq(api_key=st.secrets["groq_api_key"])
            prompt = f"""Analise esta atualização de projeto governamental:
- Marco: {marco}
- Responsável: {resp}
- Atualização: {texto}
- Nova data: {data if data else 'Sem alteração'}

Sua tarefa:
1. Classifique o risco (Baixo, Médio, Alto).
2. Identifique causas de atraso, dependências externas e se há bloqueio operacional.
3. Avalie o impacto no prazo.
4. Crie uma sugestão de resumo bem curta e direta (máximo de 2 a 3 frases) ideal para o técnico copiar e colar como registro oficial na planilha. Inicie OBRIGATORIAMENTE essa sugestão com a frase: "Eu resumiria essa atualização para a planilha como: "

Seja objetivo e use bullet points para os itens 1 a 3."""
            resposta = client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model="llama-3.1-8b-instant",
                temperature=0.3
            )
            return resposta.choices[0].message.content
        except Exception as e:
            return f"Erro na IA: {e}"

    for i, template_nome in enumerate(TEMPLATES.keys()):
        with tabs[i + 1]:
            respostas = todas_respostas[todas_respostas['template'] == template_nome] if not todas_respostas.empty else todas_respostas

            if not respostas.empty:
                col_lista, col_ia = st.columns([2, 1])
                with col_lista:
                    for _, row in respostas.iterrows():
                        concluido_label = "✅ Sim" if row.get('concluido') == "Sim" else "❌ Não"
                        with st.expander(f"📥 {row['marco']} — {row['responsavel']} | Concluída: {concluido_label}"):
                            st.markdown("**Revise ou altere as informações antes de gravar no Excel:**")

                            texto_editado = st.text_area(
                                "Atualização da equipe (editável):",
                                value=row['atualizacao'],
                                key=f"edit_txt_{row['id']}",
                                height=150
                            )

                            data_pre_preenchida = None
                            if row['nova_data'] and row['nova_data'] not in ["None", ""]:
                                try:
                                    data_pre_preenchida = datetime.strptime(row['nova_data'], "%Y-%m-%d").date()
                                except:
                                    data_pre_preenchida = None

                            data_editada = st.date_input(
                                "Nova tendência de término:",
                                value=data_pre_preenchida,
                                key=f"edit_dt_{row['id']}"
                            )

                            st.divider()

                            b1, b2 = st.columns(2)
                            if b1.button("✅ Confirmar na Planilha", key=f"conf_{row['id']}", type="primary"):
                                with st.spinner("Gravando edições no Excel..."):
                                    data_para_gravar = str(data_editada) if data_editada else ""
                                    if confirmar_na_planilha(row['hash_id'], texto_editado, data_para_gravar, row.get('template') or 'reurb'):
                                        deletar_resposta(row['id'])
                                        st.cache_data.clear()
                                        st.success("Atualização gravada com sucesso!")
                                        st.rerun()
                                    else:
                                        st.error("Erro ao gravar. Verifique se o arquivo Excel está fechado.")

                            if b2.button("🧠 Análise IA", key=f"ia_{row['id']}"):
                                st.session_state[f'ia_res_{template_nome}']   = analisar_com_ia(row['marco'], row['responsavel'], texto_editado, str(data_editada))
                                st.session_state[f'ia_marco_{template_nome}'] = row['marco']

                with col_ia:
                    st.markdown("### 🤖 AIest")
                    st.caption("Agente Auxiliar de Monitoramento Estratégico")
                    st.divider()
                    if f'ia_res_{template_nome}' in st.session_state:
                        st.markdown(f"**Análise referente a:** *{st.session_state[f'ia_marco_{template_nome}']}*")
                        st.info(st.session_state[f'ia_res_{template_nome}'])
                    else:
                        st.write("👈 Clique em **'🧠 Análise IA'** para receber o diagnóstico.")
            else:
                st.success("Nenhuma atualização pendente.")

else:
    # ---------------------------------------------------------
    # TELA DO RESPONSÁVEL (MÚLTIPLOS CAMPOS)
    # ---------------------------------------------------------
    if st.session_state.get("formulario_enviado"):
        st.image("AEST Sede (1).png", width=160)
        st.title("✅ Formulário respondido")
        st.success("Suas atualizações foram registradas com sucesso. Você já pode fechar esta página.")
        st.stop()

    st.title("📝 Atualização de Status")
    lista_hashes = hash_url.split(',')
    res = df_base[df_base["HASH_ID"].isin(lista_hashes)]

    if not res.empty:
        ja_respondido, quem_respondeu, data_resp = verificar_formulario_respondido(lista_hashes)
        if ja_respondido:
            st.warning(f"⚠️ Atualizações já enviadas por **{quem_respondeu}** em {data_resp}.")
            st.stop()

        st.info("Por favor, preencha as atualizações para os marcos abaixo:")

        with st.form("f_multi"):
            respondido_por = st.text_input("Respondido por:", placeholder="Seu nome completo", key="respondido_por")
            st.divider()

            # O 'idx' representa o número real da linha do dataframe (sempre único)
            for idx, r in res.iterrows():
                data_formatada = pd.to_datetime(r['Tendência de Término']).strftime('%d/%m/%Y') if r['Tendência de Término'] else 'Não definido'

                st.markdown(f"### 📌 {r['Marcos Críticos para Realizar as Entregas']}")
                st.write(f"**Responsável:** {r['Responsável']} | **Prazo Atual:** {data_formatada}")

                st.text_area("Atualização da equipe:", key=f"txt_{r['HASH_ID']}_{idx}")
                st.date_input("Nova tendência de término:", value=None, key=f"dt_{r['HASH_ID']}_{idx}")
                st.radio("Atividade concluída?", options=["Não", "Sim"], horizontal=True, key=f"conc_{r['HASH_ID']}_{idx}")
                st.divider()

            if st.form_submit_button("Enviar todas as atualizações", type="primary"):
                nome = st.session_state["respondido_por"].strip()
                if not nome:
                    st.error("⚠️ Informe seu nome no campo 'Respondido por' antes de enviar.")
                else:
                    enviou_algum = False
                    for idx, r in res.iterrows():
                        status    = st.session_state[f"txt_{r['HASH_ID']}_{idx}"]
                        nova_tend = st.session_state[f"dt_{r['HASH_ID']}_{idx}"]
                        concluido = st.session_state[f"conc_{r['HASH_ID']}_{idx}"]

                        if status:
                            salvar_resposta(r['HASH_ID'], r['Marcos Críticos para Realizar as Entregas'], r['Responsável'], status, nova_tend, concluido, template_ativo)
                            enviou_algum = True

                    if enviou_algum:
                        marcar_formulario_respondido(lista_hashes, nome)
                        st.session_state["formulario_enviado"] = True
                        st.rerun()
                    else:
                        st.warning("⚠️ Nenhuma atualização preenchida. Escreva em pelo menos um marco.")
    else:
        st.error("Link de atualização inválido ou atividades não encontradas.")